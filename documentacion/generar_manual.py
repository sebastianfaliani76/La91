from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from pathlib import Path

ROOT = Path(__file__).resolve().parent.parent
OUT = ROOT / 'documentacion' / 'Manual_de_Usuario_La_91_v0.16.docx'
LOGO = ROOT / 'frontend' / 'public' / 'marca' / 'logo-horizontal-claro.png'
CAPTURAS = ROOT / 'documentacion' / 'capturas_manual'
CAPTURA = CAPTURAS / '01-acceso.png'
AZUL = '003B46'; VERDE = '07575B'; CELESTE = '66A5AD'; CLARO = 'C4DFE6'; GRIS = '52656A'

doc = Document()
sec = doc.sections[0]
sec.page_width, sec.page_height = Inches(8.5), Inches(11)
sec.top_margin = sec.bottom_margin = sec.left_margin = sec.right_margin = Inches(1)
sec.header_distance = sec.footer_distance = Inches(.49)

def font(run, size=11, color=AZUL, bold=False, italic=False):
    run.font.name = 'Aptos'; run._element.get_or_add_rPr().rFonts.set(qn('w:ascii'), 'Aptos'); run._element.rPr.rFonts.set(qn('w:hAnsi'), 'Aptos')
    run.font.size = Pt(size); run.font.color.rgb = RGBColor.from_string(color); run.bold = bold; run.italic = italic
    return run

styles = doc.styles
normal = styles['Normal']; normal.font.name = 'Aptos'; normal.font.size = Pt(10.5); normal.font.color.rgb = RGBColor.from_string(AZUL)
normal.paragraph_format.space_after = Pt(6); normal.paragraph_format.line_spacing = 1.15
for name, size, before, after, color in [('Title',30,0,8,AZUL),('Subtitle',14,0,8,VERDE),('Heading 1',18,18,9,AZUL),('Heading 2',14,14,7,VERDE),('Heading 3',11.5,10,5,AZUL)]:
    s=styles[name]; s.font.name='Aptos'; s.font.size=Pt(size); s.font.bold=name!='Subtitle'; s.font.color.rgb=RGBColor.from_string(color); s.paragraph_format.space_before=Pt(before); s.paragraph_format.space_after=Pt(after); s.paragraph_format.keep_with_next=True
for name in ['List Bullet','List Number']:
    s=styles[name]; s.font.name='Aptos'; s.font.size=Pt(10.5); s.paragraph_format.left_indent=Inches(.38); s.paragraph_format.first_line_indent=Inches(-.19); s.paragraph_format.space_after=Pt(4); s.paragraph_format.line_spacing=1.15

def shade(cell, fill):
    tcPr=cell._tc.get_or_add_tcPr(); shd=tcPr.find(qn('w:shd'))
    if shd is None: shd=OxmlElement('w:shd'); tcPr.append(shd)
    shd.set(qn('w:fill'),fill)

def set_cell_margin(cell, top=90, start=120, bottom=90, end=120):
    tc=cell._tc; tcPr=tc.get_or_add_tcPr(); tcMar=tcPr.first_child_found_in('w:tcMar')
    if tcMar is None: tcMar=OxmlElement('w:tcMar'); tcPr.append(tcMar)
    for m,v in [('top',top),('start',start),('bottom',bottom),('end',end)]:
        node=tcMar.find(qn(f'w:{m}'))
        if node is None: node=OxmlElement(f'w:{m}'); tcMar.append(node)
        node.set(qn('w:w'),str(v)); node.set(qn('w:type'),'dxa')

def table(headers, rows, widths=None):
    t=doc.add_table(rows=1, cols=len(headers)); t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.autofit=False; t.style='Table Grid'
    for i,h in enumerate(headers):
        c=t.rows[0].cells[i]; c.text=''; shade(c,AZUL); set_cell_margin(c); c.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER; font(c.paragraphs[0].add_run(h),9.5,'FFFFFF',True)
    trPr=t.rows[0]._tr.get_or_add_trPr(); tblHeader=OxmlElement('w:tblHeader'); tblHeader.set(qn('w:val'),'true'); trPr.append(tblHeader)
    for row in rows:
        cells=t.add_row().cells
        for i,value in enumerate(row):
            cells[i].text=''; set_cell_margin(cells[i]); cells[i].vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER; font(cells[i].paragraphs[0].add_run(str(value)),9.2,AZUL)
    if widths:
        for row in t.rows:
            for i,w in enumerate(widths): row.cells[i].width=Inches(w)
    doc.add_paragraph().paragraph_format.space_after=Pt(2)
    return t

def note(title, text, color=CLARO):
    p=doc.add_paragraph(); p.paragraph_format.left_indent=Inches(.12); p.paragraph_format.right_indent=Inches(.12); p.paragraph_format.space_before=Pt(5); p.paragraph_format.space_after=Pt(8); p.paragraph_format.keep_together=True
    pPr=p._p.get_or_add_pPr(); shd=OxmlElement('w:shd'); shd.set(qn('w:fill'),color); pPr.append(shd)
    borders=OxmlElement('w:pBdr')
    for side in ('top','left','bottom','right'):
        border=OxmlElement(f'w:{side}'); border.set(qn('w:val'),'single'); border.set(qn('w:sz'),'4'); border.set(qn('w:space'),'5'); border.set(qn('w:color'),CELESTE); borders.append(border)
    pPr.append(borders); font(p.add_run(title+' '),10,AZUL,True); font(p.add_run(text),10,AZUL)

def bullets(items):
    for x in items: doc.add_paragraph(x,style='List Bullet')

def steps(items):
    for x in items: doc.add_paragraph(x,style='List Number')

def page(): doc.add_page_break()

def heading(title, level=1): doc.add_heading(title,level=level)

def screenshot_placeholder(modulo):
    capturas = {
        'Inicio': [('02-inicio.png', 'Tablero de inicio con el resumen operativo.')],
        'Catálogo': [('03-catalogo.png', 'Catálogo, categorías y búsqueda de productos.')],
        'Inventario': [('04-inventario.png', 'Existencias, filtros y control de stock mínimo.')],
        'Compras y proveedores': [
            ('05-proveedores.png', 'Padrón y cuentas de proveedores.'),
            ('06-compras.png', 'Órdenes de compra y sus estados.'),
        ],
        'Clientes': [('07-clientes.png', 'Padrón de clientes y cuentas corrientes.')],
        'Punto de venta': [('08-punto-venta.png', 'Punto de venta y apertura de caja requerida.')],
        'Cierre de caja': [('08-punto-venta.png', 'Acceso al punto de venta y gestión de la caja.')],
        'Gastos': [('10-gastos.png', 'Gastos, servicios y obligaciones pendientes.')],
        'Empleados': [
            ('11-empleados.png', 'Padrón de empleados y resumen de nómina.'),
            ('15-detalle-empleado.png', 'Detalle de liquidaciones, pagos y adelantos.'),
        ],
        'Tesorería': [('12-tesoreria.png', 'Cuentas, obligaciones y libro de tesorería.')],
        'Usuarios': [
            ('13-usuarios.png', 'Usuarios, roles y estados de acceso.'),
            ('14-nuevo-usuario-modal.png', 'Formulario para crear un usuario.'),
        ],
        'Reportes': [('09-reportes.png', 'Análisis comercial agrupado por circuito.')],
        'Tienda online': [('15-tienda-online.png', 'Tienda online con promociones, categorías, productos y paginación.')],
        'Compra online': [
            ('16-carrito-online.png', 'Carrito con cantidades, subtotales y acceso a la confirmación.'),
            ('17-checkout-online.png', 'Formulario de entrega, pago, cupón y confirmación del pedido.'),
        ],
        'Preparación online': [('18-preparacion-pedido-online.png', 'Detalle del pedido y control compacto de las cantidades preparadas.')],
        'Control pedido online': [('19-control-pago-pedido-online.png', 'Totales, pago aprobado y control de reintegros del pedido.')],
    }
    for archivo, descripcion in capturas.get(modulo, []):
        ruta = CAPTURAS / archivo
        if not ruta.exists():
            continue
        p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(7); p.paragraph_format.space_after=Pt(3); p.paragraph_format.keep_together=True; p.paragraph_format.keep_with_next=True
        pic=p.add_run().add_picture(str(ruta),width=Inches(6.35)); pic._inline.docPr.set('descr',descripcion)
        if archivo == '18-preparacion-pedido-online.png':
            blip_fill = pic._inline.xpath('.//pic:blipFill')[0]
            src_rect = OxmlElement('a:srcRect')
            src_rect.set('t', '43000')
            blip_fill.insert(1, src_rect)
        p=doc.add_paragraph(descripcion); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after=Pt(9); p.paragraph_format.keep_with_next=False; font(p.runs[0],8.5,GRIS,False,True)

# Encabezado y pie
h=sec.header.paragraphs[0]; h.alignment=WD_ALIGN_PARAGRAPH.LEFT; font(h.add_run('LA 91 SUPERMERCADO  |  MANUAL DE USUARIO'),8.5,VERDE,True)
f=sec.footer.paragraphs[0]; f.alignment=WD_ALIGN_PARAGRAPH.CENTER; font(f.add_run('Versión preliminar 0.16 · Agosto 2026 · Uso interno y capacitación'),8,GRIS)

# Portada editorial
doc.add_paragraph().paragraph_format.space_after=Pt(70)
if LOGO.exists():
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; pic=p.add_run().add_picture(str(LOGO),width=Inches(4.9)); pic._inline.docPr.set('descr','Logotipo de La 91 Supermercado')
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; font(p.add_run('MANUAL DE USUARIO'),12,CELESTE,True)
p=doc.add_paragraph(); p.style='Title'; p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.add_run('Sistema de Gestión\nLa 91 Supermercado')
p=doc.add_paragraph(); p.style='Subtitle'; p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.add_run('Guía operativa para administración, supervisión y caja')
doc.add_paragraph().paragraph_format.space_after=Pt(90)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; font(p.add_run('VERSIÓN PRELIMINAR 0.16'),11,VERDE,True)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; font(p.add_run('Agosto de 2026'),10,GRIS)
note('Estado del documento.', 'Esta edición incorpora capturas reales obtenidas con los datos de la prueba funcional. Se actualizará cuando cambien pantallas o circuitos.', 'EAF4F6')

page(); heading('Control del documento')
table(['Campo','Detalle'],[
 ('Documento','Manual de usuario del Sistema de Gestión La 91 Supermercado'),('Versión','0.16 – preliminar ilustrada'),('Fecha','Agosto de 2026'),('Destinatarios','Administrador, supervisor, cajero y personal autorizado'),('Objetivo','Explicar la operación habitual y los controles necesarios'),('Próxima revisión','Cuando cambien pantallas o circuitos operativos')],[1.65,4.85])
heading('Cómo utilizar este manual',2)
doc.add_paragraph('Cada capítulo describe el objetivo del módulo, el procedimiento habitual y los controles que deben respetarse. Los nombres de botones y opciones aparecen tal como se muestran en el sistema.')
note('Regla general.', 'Nunca comparta su contraseña. No elimine ni corrija movimientos financieros por fuera del procedimiento autorizado. Ante una diferencia, registre lo ocurrido y comuníquelo al supervisor.')
heading('Contenido')
contents=['1. Acceso y navegación','2. Roles y responsabilidades','3. Inicio y tablero','4. Catálogo de productos','5. Inventario','6. Proveedores y compras','7. Clientes y cuentas corrientes','8. Punto de venta','9. Cambios y devoluciones','10. Caja','11. Gastos y servicios','12. Empleados y sueldos','13. Tesorería','14. Reportes','15. Tienda online y e-commerce','16. Cierre diario y controles','17. Problemas frecuentes','18. Lista de verificación para puesta en marcha']
for item in contents:
    p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(2); font(p.add_run(item),9.5,AZUL)

page(); heading('1. Acceso y navegación')
heading('1.1 Iniciar sesión',2)
steps(['Abra el navegador e ingrese a la dirección proporcionada por el responsable del sistema.','Escriba su nombre de usuario.','Ingrese su contraseña personal.','Seleccione Ingresar.','Compruebe que en la esquina superior figure su usuario.'])
if CAPTURA.exists():
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; pic=p.add_run().add_picture(str(CAPTURA),width=Inches(5.8)); pic._inline.docPr.set('descr','Pantalla de acceso al Sistema de Gestión La 91 Supermercado')
    p=doc.add_paragraph('Pantalla de acceso al sistema.'); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; font(p.runs[0],8.5,GRIS,False,True)
heading('1.2 Cerrar sesión',2)
doc.add_paragraph('Al terminar, seleccione Cerrar sesión en la barra superior. No deje una sesión abierta en un equipo compartido.')
heading('1.3 Navegación',2)
doc.add_paragraph('El menú superior muestra únicamente los módulos habilitados para el usuario. Si una opción no aparece, consulte al administrador; no significa necesariamente que exista un error.')

heading('2. Roles y responsabilidades')
table(['Rol','Responsabilidad principal','Operaciones habituales'],[
 ('Administrador','Configuración, accesos y control general','Todas las operaciones, usuarios, roles y permisos'),('Supervisor','Dirección operativa del negocio','Operación completa, supervisión, finanzas, personal y reportes; usuarios en consulta'),('Depósito','Mercadería y abastecimiento','Existencias, movimientos, órdenes, compras y recepción'),('Cajero','Atención y manejo de su caja','Ventas, cobranzas, cambios, devoluciones y cierre de su caja')],[1.1,2.15,3.25])
heading('2.1 Cajero',2)
doc.add_paragraph('Atiende al cliente y responde por el dinero de su turno. Puede consultar productos, existencias y clientes; abrir y cerrar su caja; vender; cobrar cuentas corrientes; y registrar cambios o devoluciones autorizadas.')
bullets(['Utilizar siempre su propio usuario y una caja física disponible.','Comprobar efectivo inicial, cobros, vuelto y efectivo contado.','No anular ventas completas, modificar precios, ajustar stock ni acceder a Tesorería.','Informar al supervisor cualquier diferencia antes de cerrar.'])
heading('2.2 Depósito',2)
doc.add_paragraph('Controla la mercadería y el abastecimiento. Detecta faltantes, revisa mínimos, prepara órdenes de compra, realiza o coordina pedidos y controla la recepción total o parcial.')
bullets(['Consultar catálogo y existencias sin modificar precios.','Registrar movimientos entre ubicaciones cuando corresponda.','Preparar, enviar y recibir órdenes de compra.','No realizar ajustes manuales, pagos, ventas, sueldos ni movimientos financieros.'])
note('Compra presencial.', 'Si el responsable va al mayorista, utiliza una orden como lista de compra, conserva al mayorista como proveedor y registra la recepción al regresar. Administración registra la factura y Tesorería registra el pago desde el origen real.')
heading('2.3 Supervisor',2)
doc.add_paragraph('Administra la operación general. Puede gestionar catálogo, inventario, compras, clientes, proveedores, cajas, ventas, gastos, empleados, Tesorería y reportes. También puede anular ventas y rendir cajas pendientes.')
bullets(['Controlar diferencias, excepciones y operaciones sensibles.','Consultar usuarios para identificar responsables.','No crear, editar, reactivar ni cambiar roles o contraseñas de usuarios.','Solicitar al administrador cualquier modificación de accesos.'])
heading('2.4 Administrador',2)
doc.add_paragraph('Posee las capacidades del supervisor y además administra identidades y accesos. Es el único responsable del alta, modificación, activación y asignación de roles de los usuarios.')
bullets(['Crear usuarios con el rol mínimo necesario.','Proteger y restablecer credenciales de manera segura.','Revisar permisos y accesos periódicamente.','No compartir cuentas administrativas para tareas cotidianas.'])
screenshot_placeholder('Usuarios')
heading('2.5 Cómo colaboran los roles',2)
doc.add_paragraph('El circuito de abastecimiento recomendado es: Depósito detecta faltantes → prepara la orden → Supervisor revisa cuando corresponda → se realiza el pedido o compra presencial → Depósito recibe y controla → Administración registra la factura → Tesorería realiza el pago.')
note('Comercios pequeños.', 'Una persona puede cumplir varias funciones, pero debe respetar el circuito y utilizar el usuario autorizado. Compartir contraseñas elimina la trazabilidad de quién realizó cada operación.')
note('Separación de funciones.', 'Cada cajero debe utilizar su propio usuario y su propia caja. La caja física no debe abrirse simultáneamente desde dos sesiones.')

page(); heading('3. Inicio y tablero')
doc.add_paragraph('Inicio resume el estado operativo: ventas del día, inventario, compras, cajas abiertas, cuentas por cobrar y pagar, gastos y sueldos pendientes.')
bullets(['Revise productos bajo mínimo antes de realizar pedidos.','Controle compras demoradas.','Observe vencimientos de clientes, proveedores y servicios.','Verifique cuántas cajas permanecen abiertas.','Use Actualizar para solicitar datos recientes.'])
screenshot_placeholder('Inicio')

heading('4. Catálogo de productos')
heading('4.1 Buscar y filtrar',2)
doc.add_paragraph('La búsqueda se actualiza mientras se escribe. También puede navegar mediante los iconos de categorías y filtrar por marca.')
heading('4.2 Crear o editar un producto',2)
steps(['Seleccione Nuevo producto o Editar.','Complete nombre, categoría, marca y códigos de barra.','Ingrese precio de costo.','Defina el porcentaje de margen o el precio de venta según corresponda.','Indique stock mínimo y demás controles enteros cuando se soliciten.','Guarde y verifique el resultado.'])
note('Precios.', 'El precio mayorista importado se interpreta como precio de costo. El precio de venta puede calcularse aplicando un porcentaje sobre ese costo.')
heading('4.3 Categorías y marcas',2)
doc.add_paragraph('Las altas y modificaciones se realizan siempre en ventanas modales. Mantenga nombres claros y evite crear categorías o marcas duplicadas.')
screenshot_placeholder('Catálogo')

page(); heading('5. Inventario')
doc.add_paragraph('Inventario muestra cantidad disponible, reservada y stock mínimo. Disponible representa lo utilizable; reservado corresponde a unidades comprometidas por operaciones que todavía no deben descontarse definitivamente.')
heading('5.1 Filtros',2)
bullets(['Búsqueda por nombre o código.','Categoría y marca.','Productos por debajo del mínimo.','Productos con o sin existencia.'])
heading('5.2 Ajuste o conteo físico',2)
steps(['Busque el producto.','Abra el control de existencia.','Ingrese la cantidad física contada.','Indique el motivo de la diferencia.','Confirme y compruebe el nuevo saldo.'])
note('Control.', 'Todo ajuste debe tener una causa real: conteo, rotura, merma, corrección documentada o consumo interno. No utilice ajustes para ocultar errores de venta o compra.')
screenshot_placeholder('Inventario')

heading('6. Proveedores y compras')
heading('6.1 Proveedores',2)
doc.add_paragraph('Registre razón social, nombre comercial, CUIT, contacto y condiciones relevantes. La cuenta del proveedor muestra facturas, pagos, saldo y deuda vencida.')
heading('6.2 Orden de compra',2)
steps(['Seleccione Nueva orden.','Elija el proveedor.','Busque productos escribiendo nombre o código.','Use flechas y Enter para agregar artículos rápidamente.','Complete cantidades y costos.','Revise el total y guarde la orden.','Al recibir mercadería, confirme las cantidades realmente entregadas.'])
heading('6.3 Facturas y pagos',2)
steps(['Abra la cuenta del proveedor.','Registre la factura y su vencimiento; vincúlela a la compra si corresponde.','Registre pagos parciales o totales.','Seleccione el medio correcto.','Compruebe el saldo restante y el comprobante.'])
note('Compras realizadas personalmente.', 'Aunque el dueño o empleado retire la mercadería, la compra debe registrarse al proveedor que emitió la factura o entregó los productos.')
screenshot_placeholder('Compras y proveedores')

page(); heading('7. Clientes y cuentas corrientes')
heading('7.1 Crear un cliente',2)
steps(['Ingrese a Clientes y seleccione Nuevo cliente.','Complete nombre y documento.','Si tendrá crédito, habilite la cuenta corriente.','Defina límite de crédito y días de vencimiento.','Guarde los cambios.'])
heading('7.2 Venta a crédito y cobranzas',2)
doc.add_paragraph('En una venta puede registrarse un pago inicial y dejar el resto en cuenta corriente. El sistema controla el límite disponible. Las cobranzas posteriores reducen el saldo del cliente.')
steps(['Abra la cuenta del cliente.','Revise ventas pendientes y vencidas.','Seleccione Registrar cobranza.','Indique importe, medio y referencia.','Verifique el nuevo saldo.'])
note('Crédito responsable.', 'No habilite crédito sin autorización. Antes de vender, revise saldo, límite disponible y vencimientos anteriores.')
screenshot_placeholder('Clientes')

heading('8. Punto de venta')
heading('8.1 Antes de vender',2)
doc.add_paragraph('El usuario debe tener una caja abierta. Al abrirla, seleccione una caja disponible e ingrese el efectivo inicial contado.')
heading('8.2 Registrar una venta',2)
steps(['Escriba o escanee el código de barras. Si existe una coincidencia exacta, el producto se agrega automáticamente.','Para buscar por nombre, escriba parte de la descripción.','Use Flecha abajo para recorrer resultados y Enter para agregar.','Mantenga la lista abierta para agregar varias unidades con Enter.','Revise cantidades, precios y stock.','Seleccione cliente cuando corresponda.','Complete uno o varios medios de pago.','Confirme la venta y entregue el comprobante disponible.'])
note('Operación rápida.', 'Flecha arriba desde el primer resultado vuelve al campo de búsqueda. Al llegar al último resultado, la navegación continúa como un único control entre búsqueda y grilla.')
screenshot_placeholder('Punto de venta')

page(); heading('9. Cambios, devoluciones y anulaciones')
doc.add_paragraph('Los cambios deben partir de una venta registrada. El sistema permite devolver un artículo, entregar otro y resolver diferencias a favor del cliente o del comercio.')
steps(['Busque la venta en Historial.','Abra su detalle.','Seleccione la operación de cambio o devolución.','Indique el producto y la cantidad devuelta.','Agregue el producto de reemplazo si corresponde.','Registre el dinero devuelto o cobrado como diferencia.','Confirme mediante el modal y revise el movimiento de stock y caja.'])
note('No eliminar.', 'Las ventas y movimientos financieros no deben borrarse. Utilice anulaciones, devoluciones o contramovimientos para conservar trazabilidad.')

heading('10. Caja')
heading('10.1 Apertura',2)
steps(['Seleccione una caja que figure disponible.','Cuente el dinero físico inicial.','Ingrese el monto y confirme.'])
heading('10.2 Movimientos y cierre',2)
doc.add_paragraph('El resumen separa ventas, cobranzas, devoluciones, ingresos manuales y egresos por proveedores, gastos, sueldos y adelantos.')
steps(['Revise el resumen de la sesión.','Cuente todo el efectivo físico.','Ingrese Efectivo contado.','Compare con Efectivo esperado.','Si existe diferencia, vuelva a contar y documente la causa.','Confirme el cierre.'])
note('Diferencias.', 'Nunca modifique ventas o movimientos para forzar el cierre. La diferencia debe quedar registrada y ser revisada por el supervisor.')
screenshot_placeholder('Cierre de caja')

page(); heading('11. Gastos y servicios')
heading('11.1 Registrar una obligación',2)
steps(['Seleccione Nuevo gasto.','Indique concepto y categoría.','Asocie proveedor o beneficiario si existe.','Complete comprobante, total, emisión y vencimiento.','Si se repite, marque Recurrente y elija frecuencia.','Guarde.'])
heading('11.2 Pagar y renovar',2)
doc.add_paragraph('Los pagos pueden ser parciales. El saldo queda pendiente hasta completarse. Para un servicio recurrente, Generar próximo período crea una obligación nueva conservando sus datos básicos.')
note('Efectivo.', 'Para pagar en efectivo debe existir una caja abierta para el usuario. Ese importe reduce automáticamente el efectivo esperado de la caja.')
screenshot_placeholder('Gastos')

heading('12. Empleados, sueldos y adelantos')
heading('12.1 Legajo',2)
doc.add_paragraph('Registre datos personales, cargo, modalidad diaria/semanal/quincenal/mensual, sueldo base por período y CBU o alias.')
heading('12.2 Adelanto',2)
steps(['Abra el empleado.','Seleccione Nuevo adelanto.','Ingrese fecha, importe y medio.','Si es efectivo, confirme que la caja esté abierta.','Guarde y revise que el adelanto figure pendiente.'])
heading('12.3 Liquidación y pago',2)
steps(['Seleccione Liquidar sueldo.','Defina período desde y hasta.','Revise sueldo base.','Agregue adicionales y descuentos.','Mantenga Aplicar adelantos si deben descontarse.','Cree la liquidación y revise el neto.','Registre uno o varios pagos hasta cancelar el saldo.'])
note('Cálculo.', 'Neto = sueldo base + adicionales − descuentos − adelantos aplicados. Esta liquidación es un control administrativo y no reemplaza la documentación laboral o impositiva exigible.')
screenshot_placeholder('Empleados')

page(); heading('13. Tesorería')
doc.add_paragraph('Tesorería controla fondos fuera de las cajas operativas: efectivo general, bancos, billeteras virtuales e inversiones.')
heading('13.1 Cuentas',2)
steps(['Seleccione Nueva cuenta.','Indique nombre y tipo.','Ingrese el saldo inicial real.','Guarde y compruebe el saldo.'])
heading('13.2 Movimientos',2)
doc.add_paragraph('Registre ingresos y egresos con categoría, concepto, importe, fecha y referencia. El libro permite filtrar por cuenta y tipo.')
heading('13.3 Transferencias',2)
steps(['Seleccione Transferir.','Indique cuenta de origen y destino.','Ingrese el importe.','Agregue concepto y referencia.','Confirme. El sistema registra ambas contrapartidas.'])
note('Evitar duplicaciones.', 'Hasta completar la integración automática definitiva, verifique que un pago ya registrado en Proveedores, Gastos o Sueldos no vuelva a ingresarse manualmente como egreso sin una referencia clara.')
screenshot_placeholder('Tesorería')

heading('14. Reportes')
doc.add_paragraph('Seleccione el período y revise ventas, operaciones, ticket promedio, costo, margen bruto, crédito otorgado, deudas, gastos y sueldos.')
bullets(['Compare ventas con costo y margen.','Revise productos y categorías más vendidos.','Controle medios de pago.','Analice cuentas por cobrar y pagar.','Revise gastos y sueldos pagados y pendientes.'])
note('Alcance.', 'Los reportes son herramientas de gestión. No sustituyen libros contables, declaraciones impositivas ni documentación fiscal oficial.')
screenshot_placeholder('Reportes')

page(); heading('15. Tienda online y e-commerce')
doc.add_paragraph('El canal online utiliza el mismo catálogo, los mismos clientes, las existencias y la Tesorería del sistema de gestión. Los pedidos reservan mercadería al confirmarse y se convierten en ventas cuando son entregados.')
heading('15.1 Publicar productos',2)
steps(['Abra E-commerce y seleccione Productos.','Busque el artículo.','Utilice Publicar para mostrarlo en la tienda u Ocultar para retirarlo del canal online.','Revise precio, disponibilidad, stock de seguridad, cantidad máxima y modalidades de entrega habilitadas.'])
note('Stock online.', 'La disponibilidad se obtiene de la existencia del local menos las unidades reservadas y el stock de seguridad. No se mantiene un inventario físico separado para Internet.')
screenshot_placeholder('Tienda online')
heading('15.2 Crear y administrar promociones',2)
steps(['Abra E-commerce y seleccione Promociones.','Seleccione Nueva promoción o Editar.','Defina nombre, tipo, valor, monto mínimo y vigencia.','Elija el alcance: pedido completo, categoría o productos específicos.','Active Aplicar también en el supermercado solamente si el descuento debe utilizarse en Punto de venta.','Guarde y compruebe el canal y el estado mostrados en la lista.'])
bullets(['Solo online: se aplica únicamente en la tienda.','Online + supermercado: también se calcula automáticamente en Punto de venta.','Desactivar conserva el historial e impide nuevas aplicaciones.','Programada aún no comenzó; Vencida finalizó; Inactiva fue deshabilitada manualmente.','Los cupones y el envío gratis pertenecen exclusivamente al canal online.'])
note('Orden de descuentos.', 'Primero se aplican las promociones por producto o categoría. Después se calcula la promoción general sobre el subtotal ya rebajado. El cliente debe ver precio original, ahorro, subtotal promocionado y descuento general por separado.')
heading('15.3 Comprar desde la tienda',2)
steps(['Abra Visitar la tienda.','Busque productos por nombre o navegue por categorías.','Seleccione Agregar y abra el carrito.','Revise precio original, precio promocional, ahorro y subtotal.','Seleccione Continuar.','Complete cliente, modalidad de entrega y medio de pago.','Si posee un cupón, escríbalo y seleccione Aplicar.','Revise el resumen y el importe del botón Confirmar pedido.','Confirme y conserve el código de seguimiento.'])
note('Importe antes de confirmar.', 'El resumen debe incluir subtotal original, descuentos, envío y total a pagar. El botón Confirmar pedido debe mostrar exactamente el mismo total que se registrará.')
screenshot_placeholder('Compra online')
heading('15.4 Reserva y vencimiento',2)
doc.add_paragraph('Al crear el pedido, las unidades quedan reservadas: todavía no disminuye el stock físico. Si el pago no se registra dentro del plazo configurado, el pedido pasa a Cancelado y la reserva se libera automáticamente.')
bullets(['Un pedido cancelado no puede cobrarse: el cliente debe crear uno nuevo.','Cancelar antes de pagar no genera movimientos de Tesorería ni ventas.','El plazo de reserva se configura desde E-commerce > Configuración.'])
heading('15.5 Cobrar y preparar el pedido',2)
steps(['Abra E-commerce y seleccione Pedidos.','Abra Ver y controle productos, precios, descuentos, envío y total.','Seleccione Registrar pago.','Elija cuenta de Tesorería, medio, importe, comisión y referencia.','Verifique que el pago quede Aprobado y el pedido Confirmado.','Seleccione Preparar y luego Preparación completa.','Controle que el pedido quede Listo.'])
note('Tesorería.', 'Registrar el pago genera un único ingreso en la cuenta seleccionada. Preparar o entregar no debe generar un segundo ingreso.')
screenshot_placeholder('Preparación online')
heading('15.6 Entregar y registrar la venta',2)
steps(['Con el pedido Listo y el pago aprobado, seleccione Entregar y registrar venta.','Compruebe que el pedido quede Entregado.','Revise que el stock físico disminuya y la reserva vuelva a cero.','Abra Punto de venta y seleccione Historial para consultar la venta con canal e-commerce.'])
bullets(['La venta online aparece como Venta online y no requiere una caja física.','Los reportes incorporan la venta, el costo y el margen.','El movimiento de Tesorería proviene del cobro, no de la entrega.'])
note('Circuito validado.', 'La prueba funcional confirmó el pedido Entregado, la reserva en cero, el descuento correcto del stock físico y la venta online visible en Punto de venta > Historial. Tesorería conservó un solo ingreso: el registrado al cobrar el pedido.')
screenshot_placeholder('Control pedido online')
heading('15.7 Cancelar un pedido y realizar el reembolso total',2)
steps(['Abra E-commerce > Pedidos y seleccione Ver en el pedido.','Compruebe los productos, el total pagado y la cuenta utilizada para el cobro.','Seleccione Cancelar y confirme el motivo.','Registre el reembolso desde la cuenta de Tesorería correspondiente.','Compruebe que el pedido figure Cancelado y que ya no permita registrar nuevos pagos.','Revise que la reserva de todos los productos haya sido liberada.','Controle en Tesorería el egreso por el importe reembolsado.'])
note('Control del reintegro.', 'El detalle debe separar el total pagado, el total reembolsado y cualquier saldo pendiente de reintegro. Cancelar libera la reserva, pero el movimiento de dinero se registra al confirmar el reembolso desde la cuenta seleccionada.')
heading('15.8 Devolución parcial por producto',2)
doc.add_paragraph('La devolución parcial se utiliza cuando el cliente devuelve uno o varios artículos sin cancelar el resto del pedido. El pedido conserva su operación válida y solamente se revierten las unidades seleccionadas.')
steps(['Abra el detalle del pedido pagado y seleccione Devolución parcial.','Marque el producto y la cantidad que devuelve el cliente.','Ingrese el motivo de la devolución.','Seleccione la cuenta de Tesorería desde la cual se realizará el reintegro.','Revise el importe calculado y confirme mediante el modal.','Compruebe que el pedido continúe activo y conserve los productos no devueltos.','Verifique que solamente las unidades devueltas regresen al stock.','Controle el egreso correspondiente en el libro de Tesorería y el historial dentro del pedido.'])
bullets(['El reintegro se calcula utilizando el precio promocional efectivamente pagado, no el precio original sin descuento.','No se puede devolver una cantidad superior a la comprada ni volver a devolver unidades ya reintegradas.','Una devolución parcial no cancela automáticamente todo el pedido.','El detalle del pedido conserva producto, cantidad, motivo, cuenta e importe para auditoría.'])
note('Resultado validado.', 'Durante la prueba funcional se confirmó que el pedido permaneció activo, el stock recuperó solamente la cantidad devuelta y Tesorería registró únicamente el importe reintegrado.')

page(); heading('16. Cierre diario y controles')
heading('16.1 Cajero',2)
bullets(['Finalizar ventas pendientes.','Revisar devoluciones y diferencias.','Contar efectivo.','Cerrar su caja.','Cerrar sesión.'])
heading('16.2 Supervisor o administrador',2)
bullets(['Verificar que no queden cajas abiertas sin responsable.','Revisar diferencias de cierre.','Controlar ventas y cobranzas del día.','Revisar stock bajo mínimo.','Controlar vencimientos próximos.','Confirmar movimientos relevantes de Tesorería.','Verificar que la copia de seguridad se haya completado cuando esté configurada.'])
table(['Control','Frecuencia','Responsable'],[
 ('Cierre de cada caja','Diaria / por turno','Cajero y supervisor'),('Stock bajo mínimo','Diaria','Compras o supervisor'),('Cuentas vencidas','Diaria','Administración'),('Tesorería y bancos','Diaria','Administrador'),('Conteo físico selectivo','Semanal','Supervisor'),('Respaldo restaurable','Según política definida','Responsable técnico')],[2.7,1.4,2.4])

heading('17. Problemas frecuentes')
table(['Situación','Qué revisar'],[
 ('No aparece un módulo','Permisos del usuario y necesidad de volver a iniciar sesión.'),('No encuentra un producto','Código, nombre, estado activo, precio de venta y existencia.'),('No permite pagar en efectivo','El usuario debe tener una caja abierta.'),('No permite abrir una caja','Puede estar ocupada por otra sesión.'),('El saldo no coincide','Revise pagos parciales, devoluciones, cobranzas y filtros.'),('La venta quedó a crédito','Revise cliente seleccionado, pagos cargados y saldo pendiente.'),('Hay diferencia al cerrar','Vuelva a contar y revise devoluciones, gastos, sueldos, adelantos y movimientos manuales.'),('Pedido online cancelado','Venció el plazo de reserva. El stock fue liberado y debe generarse un pedido nuevo.'),('El total online no coincide','Revise promociones, cupón, modalidad de entrega y costo de zona antes de confirmar.'),('El sistema muestra un error','Copie el mensaje, anote la operación y comuníquelo sin repetirla varias veces.')],[2.2,4.3])

page(); heading('18. Lista de verificación para puesta en marcha')
for x in ['Usuarios y roles revisados','Cajas físicas identificadas','Productos, precios y códigos verificados','Stock inicial controlado','Proveedores cargados','Clientes con crédito autorizados','Empleados y modalidades revisados','Cuentas de Tesorería conciliadas','Prueba de compra realizada','Prueba de venta y devolución realizada','Prueba de cuenta corriente y cobranza realizada','Prueba de gasto, sueldo y adelanto realizada','Promociones online y locales verificadas','Pedido online, reserva, cobro, preparación y entrega comprobados','Venta online visible en historial y reportes','Cierre de caja comprobado','Reportes revisados','Procedimiento de respaldo y restauración probado']:
    doc.add_paragraph('☐ '+x)
note('Criterio de salida.', 'La versión 1.0 del manual se emitirá después de la puesta en marcha controlada y de incorporar cualquier corrección detectada durante el uso real.')
heading('Registro de incidencias de la prueba',2)
table(['Fecha','Módulo','Situación observada','Prioridad','Resolución'],[('','','','',''),('','','','',''),('','','','',''),('','','','','')],[.8,1.05,2.75,.8,1.1])

doc.core_properties.title='Manual de Usuario - Sistema de Gestión La 91 Supermercado'
doc.core_properties.subject='Guía operativa ilustrada, versión 0.16'
doc.core_properties.author='La 91 Supermercado'
doc.core_properties.keywords='supermercado, manual, usuario, caja, inventario, tesorería'
OUT.parent.mkdir(parents=True,exist_ok=True)
doc.save(OUT)
print(OUT)
